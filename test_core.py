# -*- coding: utf-8 -*-
"""核心逻辑单元测试（不导入 GUI 模块）"""

from replacer import load_rules, replace_in_docx, backup_file, restore_file, has_backup
from docx import Document

RULES_FILE = r'c:\code\替换\test_data\test_rules.xlsx'
DOC_FILE = r'c:\code\替换\test_data\test_doc.docx'

# 1. 测试规则加载
print('=== 测试规则加载 ===')
rules = load_rules(RULES_FILE)
for old, new in rules:
    print(f'  {old} -> {new}')
assert len(rules) == 5, f'期望 5 条规则，实际 {len(rules)} 条'
print(f'OK: 加载了 {len(rules)} 条规则\n')

# 2. 测试备份
print('=== 测试备份 ===')
backup_path = backup_file(DOC_FILE)
print(f'  备份路径: {backup_path}')
assert has_backup(DOC_FILE), '备份应该存在'
print('OK: 备份成功\n')

# 3. 测试替换
print('=== 测试替换 ===')
result = replace_in_docx(DOC_FILE, rules)
print(f'  替换结果: {result}')
assert result['total_replacements'] > 0, '应该有替换发生'

# 验证替换后内容
doc = Document(DOC_FILE)
all_text = ' '.join(p.text for p in doc.paragraphs)
print(f'  替换后文本: {all_text}')
assert '香蕉' in all_text, '苹果应该被替换为香蕉'
assert '蓝色' in all_text, '红色应该被替换为蓝色'
assert '上海' in all_text, '北京应该被替换为上海'
print('OK: 替换成功\n')

# 4. 测试还原
print('=== 测试还原 ===')
ok = restore_file(DOC_FILE)
assert ok, '还原应该成功'
doc2 = Document(DOC_FILE)
all_text2 = ' '.join(p.text for p in doc2.paragraphs)
print(f'  还原后文本: {all_text2}')
assert '苹果' in all_text2, '还原后应包含苹果'
assert '红色' in all_text2, '还原后应包含红色'
assert '北京' in all_text2, '还原后应包含北京'
print('OK: 还原成功\n')

print('=============================')
print('所有单元测试通过！')
